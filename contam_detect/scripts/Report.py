# -*- coding: utf-8 -*-
# @Time    : 2022/12/27 19:21
# @Author  : zzk
# @File    : Report.py
# Description:
import datetime

import docx
from docx.shared import RGBColor, Pt, Cm
import nw as NW


def sgRNAReport(doc, fa_file, seq_20nt, allData, dic_basic20):
    p0 = doc.add_paragraph()
    f = open(fa_file)
    faName = f.readline().replace("\n", "")
    ReferenceSeq = ''.join(f.readlines()).replace('\n', '')
    p01 = p0.add_run(f"{datetime.datetime.now().strftime('%Y-%m-%d')}\n{faName}\n\n")
    p01.font.size = Pt(15)
    p01.font.bold = True
    p02 = p0.add_run(f"参考序列：\n{ReferenceSeq}\n\n参考sgRNA：{seq_20nt}\n\n")
    p03 = p0.add_run("污染认定：\n")
    p03.font.size = Pt(15)
    p03.font.bold = True
    allConsenCont = ""
    with open('./racon/consensus_racon.fa') as f:
        cont1 = f.readline()
        while cont1:
            cont2 = f.readline().replace("\n", "")
            rate = round(int(cont1.replace("\n", "").split("ubs=")[1]) / allData, 3)
            name, num = cont1.replace('\n', '').split('ubs=')[0], int(cont1.replace("\n", "").split("ubs=")[1])
            allConsenCont += f"{name}\n判定污染 {num} reads,污染率为:{rate}.\n污染共识序列：\n{cont2}\n\nsgRNA为：\n{dic_basic20[f'{name}']}\n\n"
            '''
            缺一个sgRNA比对。
            '''
            sgRNA, reference_sgRNA = dic_basic20[f'{name}'], dic_basic20[f'reference-{name}']
            if len(sgRNA) - len(reference_sgRNA) > 0:
                reference_sgRNA = reference_sgRNA + "N" * (len(sgRNA) - len(reference_sgRNA))
            else:
                sgRNA = sgRNA + "N" * (len(reference_sgRNA) - len(sgRNA))
            clustal_result = NW.dynamic_edit(reference_sgRNA, sgRNA)
            allConsenCont += f"R:{clustal_result[0]}\n  {'|' * len(clustal_result[0])}\nS:{clustal_result[1]}\n\n"
            cont1 = f.readline()
    p1 = doc.add_paragraph()
    p11 = p1.add_run(allConsenCont)
    p11.font.name = "MS Mincho"
    doc.save(f"{faName.strip('>')}sgRNA-polutionReportNEW.docx")


def commonReport(doc, fa_file, allData, top_n_file):
    p0 = doc.add_paragraph()
    f = open(fa_file)
    faName = f.readline().replace("\n", "")
    ReferenceSeq = ''.join(f.readlines()).replace('\n', '')
    p01 = p0.add_run(
        f"{datetime.datetime.now().strftime('%Y-%m-%d')}\n{faName}\n参考序列：\n{ReferenceSeq}\n")
    allConsenCont = ""
    with open('./racon/consensus_racon.fa') as f:
        cont1 = f.readline()
        print(cont1)
        print(top_n_file)
        while cont1:
            if any(item.split("/")[-1] in cont1 for item in top_n_file):
                cont2 = f.readline().replace("\n", "")
                rate = round(int(cont1.replace("\n", "").split("ubs=")[1]) / allData, 3)
                '''
                >umi1_S1361-cluster-top-1_bins;ubs=184
                '''
                name, num = cont1.replace('\n', '').split('ubs=')[0], int(cont1.replace("\n", "").split("ubs=")[1])
                if fa_file[:-6].split("/")[-1] in name:
                    allConsenCont += f"{name}\n总数据量为 {allData} 判定污染{num}reads{name}\n该污染率为:{rate}.\n共识序列为：\n{cont2}\n"
                    # break
            '''
            加上与参考序列比对
            '''
            cont1 = f.readline()
    p1 = doc.add_paragraph()
    p11 = p1.add_run(allConsenCont)
    p11.font.name = "MS Mincho"
    doc.save(f"{faName.strip('>')}common-polutionReport.docx")


def consensusReport(allData, dic_basic20):
    pass
    # return allConsenCont


def run(doc, allData, dic_basic20, fa_file, seq_20nt, top_n_file):
    if seq_20nt:
        sgRNAReport(doc, fa_file, seq_20nt, allData, dic_basic20)
    else:
        docfile = commonReport(doc, fa_file, allData, top_n_file)
        return docfile
