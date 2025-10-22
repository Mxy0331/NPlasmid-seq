#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
NPlasmid-seq HTML 报告生成器（相对路径版，导航美化，BLAST对齐+横向滚动，Reads精准提取，污染报告表格）
更新：Length Distribution 一律把 PDF 转成 PNG 再显示（优先 pdftoppm，其次 ImageMagick convert）。
"""

import argparse
import os
import re
import html
import subprocess
import shutil
from pathlib import Path

# 可选依赖：python-docx，用于解析 Word
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


# ------------------------------ 通用工具 ------------------------------

def esc(s):
    return html.escape(s, quote=False) if s else ""


def _first_number_after_token(parts, token_idx):
    for j in range(token_idx + 1, len(parts)):
        m = re.search(r"\d[\d,\.]*", parts[j])
        if m:
            try:
                return int(float(re.sub(r"[^\d\.]", "", m.group(0))))
            except Exception:
                try:
                    return int(re.sub(r"[^\d]", "", m.group(0)))
                except Exception:
                    continue
    return None


def _first_number_after_pos(raw_line: str, pos: int):
    m = re.search(r"\d[\d,\.]*", raw_line[pos:])
    if not m:
        return None
    token = m.group(0)
    try:
        return int(float(re.sub(r"[^\d\.]", "", token)))
    except Exception:
        try:
            return int(re.sub(r"[^\d]", "", token))
        except Exception:
            return None


def wrap_seq_multiline(seq: str, width: int = 70) -> str:
    seq = re.sub(r"\s+", "", seq)
    return "\n".join(seq[i:i+width] for i in range(0, len(seq), width))


# ------------------------------ Reads 读取 ------------------------------

def read_demultiplex_stats(stats_path):
    stats = {}
    if not os.path.isfile(stats_path):
        return stats
    with open(stats_path, "r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            ln = raw.strip()
            if not ln or ln.startswith("#"):
                continue
            parts = [p for p in re.split(r"[\t, ]+", ln) if p != ""]
            if len(parts) < 2:
                continue
            sample_name = parts[0]
            reads = _first_number_after_token(parts, 0)
            if reads is not None:
                stats[sample_name] = reads
    return stats


def find_reads_after_sample(stats_path: str, sample: str):
    if not os.path.isfile(stats_path):
        return None
    with open(stats_path, "r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            ln = raw.rstrip("\n")
            if not ln or ln.lstrip().startswith("#"):
                continue
            parts = [p for p in re.split(r"[\t, ]+", ln) if p != ""]
            for idx, tok in enumerate(parts):
                if tok == sample:
                    num = _first_number_after_token(parts, idx)
                    if num is not None:
                        return num
            m = re.search(re.escape(sample), ln)
            if m:
                num2 = _first_number_after_pos(ln, m.end())
                if num2 is not None:
                    return num2
    return None


# ------------------------------ BLAST 解析与对齐 ------------------------------

def _format_blast_triplets_align(raw_text: str) -> str:
    lines = raw_text.splitlines()
    out = []
    i = 0
    rx_s = re.compile(r"^\s*Sbjct\s+(\d+)\s+([A-Za-z\-]+)\s+(\d+)\s*$")
    rx_q = re.compile(r"^\s*Query\s+(\d+)\s+([A-Za-z\-]+)\s+(\d+)\s*$")

    LABEL_W, START_W, END_W = 6, 7, 7

    def fmt_line(label, start, seq, end, seqw):
        return f"{label:<{LABEL_W}} {int(start):>{START_W}}  {seq:<{seqw}}  {int(end):>{END_W}}"

    while i < len(lines):
        m_s = rx_s.match(lines[i])
        if m_s and i + 2 < len(lines):
            mid = lines[i + 1]
            m_q = rx_q.match(lines[i + 2])
            if m_q:
                s_start, s_seq, s_end = m_s.groups()
                q_start, q_seq, q_end = m_q.groups()
                seqw = max(len(s_seq), len(q_seq))
                out.append(fmt_line("Sbjct", s_start, s_seq, s_end, seqw))
                prefix = f"{'':<{LABEL_W}} {'':>{START_W}}  "
                out.append(prefix + mid.lstrip())
                out.append(fmt_line("Query", q_start, q_seq, q_end, seqw))
                i += 3
                continue
        out.append(lines[i])
        i += 1
    return "\n".join(out)


def extract_blast_blocks_from_docx(docx_path, sample_name=None):
    if not HAS_DOCX or not os.path.isfile(docx_path):
        return (None, None)
    try:
        doc = Document(docx_path)
    except Exception:
        return (None, None)

    lines = [p.text.rstrip() for p in doc.paragraphs]

    def find_indices(pattern):
        rx = re.compile(pattern)
        return [i for i, ln in enumerate(lines) if rx.search(ln)]

    sample_idxs = find_indices(re.escape(sample_name)) if sample_name else []
    cons_idxs   = find_indices(r"^Consensus\s*:\s*")
    blast_idxs  = find_indices(r"BLASTN\s+2\.9\.0\+")

    def nearest(starts, anchors):
        if not starts:
            return None
        if not anchors:
            return starts[0]
        best, bestd = starts[0], 10**9
        for s in starts:
            for a in anchors:
                d = abs(s - a)
                if d < bestd:
                    best, bestd = s, d
        return best

    cons_start  = nearest(cons_idxs, sample_idxs)
    blast_start = blast_idxs[0] if blast_idxs else None

    consensus_text = None
    if cons_start is not None:
        block = []
        empty_streak = 0
        for j in range(cons_start, len(lines)):
            ln = lines[j]
            block.append(ln)
            if j > cons_start and re.match(r"^(BLASTN\s+2\.9\.0\+|Plasmid\s+Name\s*:|Score\s*=|Strand=|Identities\s*=|Consensus\s*:)", ln):
                block.pop()
                break
            if ln.strip() == "":
                empty_streak += 1
                if empty_streak >= 2:
                    break
            else:
                empty_streak = 0
        block_text = "\n".join(block)
        seqs = re.findall(r"[ACGTUNacgtun]+", block_text)
        seq_joined = "".join(seqs)
        if seq_joined:
            consensus_text = "Consensus:\n" + wrap_seq_multiline(seq_joined, width=70)
        else:
            consensus_text = block_text

    blast_text = None
    if blast_start is not None:
        rest = "\n".join(lines[blast_start:]).strip()
        blast_text = _format_blast_triplets_align(rest)

    return (consensus_text, blast_text)


# ------------------------------ PDF → PNG ------------------------------

def pdf_to_png(pdf_path: Path, out_png: Path, dpi: int = 220) -> bool:
    """
    把 pdf_path 的第一页转为 out_png（.png）。
    优先使用 pdftoppm（poppler-utils），否则尝试 ImageMagick convert。
    返回 True/False 表示是否成功生成。
    """
    try:
        pdf_path = Path(pdf_path)
        out_png = Path(out_png)
        out_png.parent.mkdir(parents=True, exist_ok=True)

        # 1) pdftoppm
        pdftoppm = shutil.which("pdftoppm")
        if pdftoppm:
            prefix = out_png.with_suffix("")
            cmd = [pdftoppm, "-singlefile", "-png", "-r", str(dpi), str(pdf_path), str(prefix)]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return out_png.is_file()

        # 2) ImageMagick convert
        convert = shutil.which("convert")
        if convert:
            cmd = [convert, "-density", str(dpi), f"{str(pdf_path)}[0]", "-quality", "92", str(out_png)]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return out_png.is_file()

        return False
    except Exception:
        return False


# ------------------------------ HTML 渲染 ------------------------------

def html_header():
    return """<!DOCTYPE html>
<html lang="zh">
<head>
<meta charset="utf-8"/>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<title>NPlasmid-seq Report</title>
<style>
:root {
  --maxw: 820px;
  --figw: 760px;
  --caption-color: #555;
  --muted: #888;
  --border: #e5e7eb;
  --navw: 340px;  /* 导航栏加宽 */
  --scrollh: 380px;
  --font: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, "Noto Sans CJK SC", "Noto Sans CJK", "PingFang SC", "Microsoft Yahei", sans-serif;
  /* 新增：标题字体栈，优先包含罗马数字字形的等权重字体 */
  --heading-font: "Segoe UI", "Segoe UI Symbol", "Arial Unicode MS", "Noto Sans Symbols 2", "Helvetica Neue", Arial, Roboto, "Noto Sans", "Noto Sans CJK SC", "PingFang SC", "Microsoft Yahei", sans-serif;
  --mono: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", "Courier New", monospace;
  --small: 14px;
  --body: 15px;
  --h1: 26px;   /* 更大主标题 */
  --h2: 22px;   /* 四个大标题更大 */
  --h3: 18px;   /* 样本名略小于大标题 */
}
* { box-sizing: border-box; }
html { scroll-behavior: smooth; }
body { margin:0; font-family: var(--font); font-size: var(--body); color:#111; background:#fff; }
a { color: #2563eb; text-decoration: none; }
a:hover { text-decoration: underline; }
.layout {
  display: grid;
  grid-template-columns: var(--navw) 1fr;
  min-height: 100vh;
}
nav {
  position: sticky; top:0;
  height: 100vh; overflow: auto;
  border-right: 1px solid var(--border);
  padding: 18px 14px;
  background: linear-gradient(180deg, #fafafa 0%, #ffffff 100%);
}
nav h2 { font-size: 24px; margin: 6px 8px 12px; color: #111; font-weight: 700; } /* Navigation 最大 */
nav ul { list-style: none; padding-left: 0; margin: 0; }
nav li { margin: 6px 0; }
nav a {
  display: block; padding: 8px 10px; border-radius: 10px; color: #374151;
  transition: background .15s ease, transform .05s ease;
  font-weight: 600;
  font-family: var(--font);
}
nav a:hover { background: #eef2ff; }
/* 四个大标题在导航中的样式：用同一标题字体栈，保证罗马数字与英文同粗 */
nav a.nav-major { font-size: 16px; font-weight: 600; font-family: var(--heading-font); }
/* 样本名在导航中更小些 */
nav a.sample-link { font-size: 14px; font-weight: 500; padding-left: 14px; }
nav .divider { height: 1px; background: var(--border); margin: 10px 6px; }

main { padding: 28px 32px 40px; }
.container { max-width: var(--maxw); margin: 0 auto; padding: 0 12px; }
/* 页面内标题也使用同一标题字体栈 */
h1 { font-size: var(--h1); text-align: center; margin: 10px 0 24px; font-weight: 700; font-family: var(--heading-font); }
h2 { font-size: var(--h2); margin: 28px 0 12px; font-weight: 600; font-family: var(--heading-font); }
h3 { font-size: var(--h3); margin: 18px 0 10px; font-weight: 600; } /* 样本名保持正常字体栈 */

p { line-height: 1.6; }
p.justify { text-align: justify; text-align-last: left; }
.section-block { margin: 22px 0 26px; }
/* 四个大块之间增加更多留白 */
.section-gap { margin: 30px 0 34px; }

.figure { width: 100%; max-width: var(--figw); margin: 12px auto 16px; text-align: center; }
.figure img { display: block; margin: 0 auto; max-width: 100%; border: 1px solid var(--border); }
#workflow .figure { overflow: hidden; }                 /* 防止溢出可见 */
#workflow .figure img { border: none !important; }     /* 保险：去掉任何 CSS 边框 */
#workflow .figure img.trim-edge { clip-path: inset(2px); } 

/* 图注两端对齐但不拉伸最后一行；tight 更紧凑 */
.caption {
  font-size: var(--small);
  color: var(--caption-color);
  text-align: justify;
  text-align-last: left;
  width: 100%; max-width: var(--figw); margin: 6px auto 2px;
}
.caption.tight { line-height: 1.35; margin: 4px auto 6px; }

.tablewrap { width: 100%; max-width: var(--figw); margin: 12px auto 10px; border: 1px solid #bcd4f6; border-radius: 8px; overflow: hidden; background: #fff; }
table { border-collapse: collapse; width: 100%; }
th, td { border: 1px solid #bcd4f6; padding: 9px 10px; font-size: var(--small); vertical-align: middle; text-align: center; }
th { background: #e8f1fe; font-weight: 700; }
.muted { color: var(--muted); }
.hr { height: 1px; background: var(--border); margin: 10px 0; }

.mono-seq { font-family: var(--mono); white-space: pre-wrap; word-break: break-word; font-size: 13px; line-height: 1.35; }

.flex2 { display: flex; gap: 18px; align-items: stretch; justify-content: space-between; margin: 8px 0 6px; }
.flex2 .box {
  flex: 1 1 0;
  border: 1px solid var(--border);
  border-radius: 8px;
  max-height: var(--scrollh);
  overflow: auto;
  overflow-x: auto;
  background: #fff;
}
.box pre {
  font-family: var(--mono);
  font-size: 13px;
  line-height: 1.35;
  padding: 10px 12px;
  margin: 0;
  white-space: pre;
  display: inline-block;
  min-width: max-content;
}

.badge { display: inline-block; font-size: 12px; padding: 2px 8px; border:1px solid var(--border); border-radius: 999px; color:#444; margin: 2px 6px 6px 0; background: #fff; }
.warning { font-size: var(--small); color:#b45309; background: #fff7ed; border:1px solid #fde68a; padding: 8px 10px; border-radius: 8px; margin: 10px auto; max-width: var(--figw); }
.section-anchor { scroll-margin-top: 14px; }

.figure img.pie-small { max-width: 70%; }
.scrollbox {
  max-height: 220px;
  overflow: auto;
  white-space: pre;
  border: 1px solid #ddd;
  padding: 8px;
  background: #fafafa;
  font-family: var(--mono);
}
</style>
</head>
<body>
<div class="layout">
"""


def html_nav(sample_list):
    nav_html = [
        "<nav>",
        "<h2>Navigation</h2>",
        "<ul>",
        '<li><a class="nav-major" href="#home">Home</a></li>',
        '<li><a class="nav-major" href="#intro">Ⅰ. Introduction</a></li>',
        '<li><a class="nav-major" href="#workflow">Ⅱ. Experimental Workflow</a></li>',
        '<li><a class="nav-major" href="#seqres">Ⅲ. Results Overview</a></li>',
        '<li><a class="nav-major" href="#samples">Ⅳ. Samples</a></li>',
        "</ul>",
        '<div class="divider"></div>',
        "<ul>"
    ]
    if sample_list:
        for s in sample_list:
            nav_html.append(f'<li><a class="sample-link" href="#sample-{esc(s)}">{esc(s)}</a></li>')
    nav_html += ["</ul>", "</nav>"]
    return "\n".join(nav_html)


def html_main_header():
    title = "NPlasmid-seq Plasmid Quality Control and Contamination Detection Report"
    return f"""<main>
  <div class="container">
    <h1 id="home" class="section-anchor">{esc(title)}</h1>
"""


def section_intro():
    return f"""
    <div class="section-block section-gap">
      <h2 id="intro" class="section-anchor">Ⅰ. Introduction</h2>
      <p class="justify">
      NPlasmid-seq is a fast, barcode-free workflow for plasmid QC. Plasmids are Cas9-linearized, pooled, and sequenced on ONT; reads are demultiplexed by plasmid-specific end-signatures instead of barcodes. This report summarizes per-plasmid quality control (QC) results generated by the NPlasmid-seq workflow. An index links to each sample, followed by standardized panels: read counts, IGV coverage views, length-distribution plots, BLAST-based consensus alignment, and contamination assessment. Together, these views let you quickly verify sequence integrity, inspect coverage and size profiles, confirm the polished consensus against its reference, and screen for cross-sample or external contaminants.
      </p>
    </div>
"""


def section_workflow():
    workflow_png = Path("workflow.png")
    exists = workflow_png.is_file()
    img_html = (
        f'<img class="trim-edge" src="{esc(str(workflow_png))}" alt="workflow"/>'
        if exists else
        '<div class="warning">找不到工作流图片：workflow.png</div>'
    )
    return f"""
    <div class="section-block section-gap">
      <h2 id="workflow" class="section-anchor">Ⅱ. Experimental Workflow</h2>
      <div class="figure">
        {img_html}
      </div>
      <div style="height:14px"></div>
    </div>
"""


def html_results_overview(sample_list):
    badges = " ".join([f'<span class="badge">{esc(s)}</span>' for s in sample_list])
    return f"""
    <div class="section-block section-gap">
      <h2 id="seqres" class="section-anchor">Ⅲ. Results Overview</h2>
      <p>Samples: {badges}</p>
      <div class="caption tight">Note: All plasmids’ names in the same run are shown here.</div>
      <div class="hr"></div>
    </div>
"""


def render_reads_table(sample, reads_map, stats_path):
    reads = reads_map.get(sample)
    if reads is None:
        reads = find_reads_after_sample(stats_path, sample)

    table_html = [
        '<div class="section-block">',
        "<h3>1. Reads Number</h3>",
        '<div class="tablewrap">',
        "<table>",
        "<thead><tr><th>Sample No.</th><th>Read Count</th></tr></thead>",
        "<tbody>"
    ]
    if reads is not None:
        table_html.append(f"<tr><td>{esc(sample)}</td><td>{reads}</td></tr>")
    else:
        table_html.append(f'<tr><td>{esc(sample)}</td><td class="muted">NA</td></tr>')
    table_html.append("</tbody></table></div>")
    table_html.append('<div class="caption tight">Note: the read count includes only reads with both ends intact.</div>')
    table_html.append("</div>")
    return "\n".join(table_html)


def render_igv_png_and_pdf(sample):
    # IGV PNG
    png_path = Path("IGV") / f"{sample}.sorted_by_strand.png"
    if png_path.is_file():
        igv_html = f"""
        <div class="figure">
          <img src="{esc(str(png_path))}" alt="IGV {esc(sample)}"/>
        </div>
        <div class="caption tight">
          <strong>IGV Visualization</strong> — Reads were aligned to the intended plasmid reference with minimap2 to generate a BAM file, which was loaded in IGV v2.19.2. The panel shows (i) a coverage track reporting per-position sequencing depth, and (ii) read-level alignments with base calls. Per-base color changes mark single-nucleotide substitutions; short insertions/deletions are rendered as small gaps or insertion markers within reads. Together, this view allows direct inspection of depth, large deletions/insertions, and small variants (SNVs/indels) along the entire plasmid.
        </div>
        """
    else:
        igv_html = f'<div class="warning">找不到 IGV PNG：{esc(str(png_path))}</div>'

    # Length Distribution：PDF -> PNG
    pdf_path = Path("Demultiplexed") / f"{sample}.pdf"
    png_alt = Path("Demultiplexed") / f"{sample}.pdf.png"

    if png_alt.is_file():
        length_block = f"""
        <h3>3. Length Distribution</h3>
        <div class="figure">
          <img src="{esc(str(png_alt))}" alt="Length histogram {esc(sample)}"/>
        </div>
        <div class="caption tight">
          <strong>Length Distribution</strong> — Read-length histogram for the sample, with a vertical line at the expected plasmid size. A single dominant peak centered on the expected length indicates clean prep and successful linearization. A left or right shoulder implies fragmented molecules. Reported summary stats (median, IQR, % within ±10% of expected length) help quantify run quality and library integrity.
        </div>
        """
    elif pdf_path.is_file():
        ok = pdf_to_png(pdf_path, png_alt, dpi=220)
        if ok and png_alt.is_file():
            length_block = f"""
            <h3>3. Length Distribution</h3>
            <div class="figure">
              <img src="{esc(str(png_alt))}" alt="Length histogram {esc(sample)}"/>
            </div>
            <div class="caption tight">
              <strong>Length Distribution</strong> — Read-length histogram for the sample, with a vertical line at the expected plasmid size. A single dominant peak centered on the expected length indicates clean prep and successful linearization. A left or right shoulder implies fragmented molecules. Reported summary stats (median, IQR, % within ±10% of expected length) help quantify run quality and library integrity.
            </div>
            """
        else:
            length_block = f'<div class="warning">无法将 PDF 转成 PNG（请安装 poppler-utils 或 imagemagick）。文件：{esc(str(pdf_path))}</div>'
    else:
        length_block = f'<div class="warning">找不到长度分布 PDF：{esc(str(pdf_path))}</div>'

    return f"""
    <div class="section-block">
      <h3>2. IGV Visualization</h3>
      {igv_html}
      {length_block}
    </div>
    """


def render_blast(sample):
    docx_path = Path("BLAST") / "doc_blast_result" / f"{sample}.docx"
    consensus_text = blast_text = None
    warn_msgs = []

    if not HAS_DOCX:
        warn_msgs.append("未安装 python-docx，无法解析 docx（请先 pip install python-docx）")
    elif not docx_path.is_file():
        warn_msgs.append(f"找不到 BLAST 文档：{docx_path}")
    else:
        c, b = extract_blast_blocks_from_docx(str(docx_path), sample_name=sample)
        consensus_text, blast_text = c, b
        if not (c or b):
            warn_msgs.append("未能从文档中解析到 Consensus/BLAST 区块（请检查文档内容格式）")

    left_box  = f'<div class="box"><pre>{esc(consensus_text) if consensus_text else "N/A"}</pre></div>'
    right_box = f'<div class="box"><pre>{esc(blast_text) if blast_text else "N/A"}</pre></div>'

    caption_html = (
        '<div class="caption tight">'
        '<strong>Consensus Sequence and BLAST Results</strong> — <strong>Left:</strong> The consensus sequence is generated from a random subset of 200 reads polished with Racon×3 followed by Medaka (SUP), reaching an accuracy of ~99.93% (Q33). '
        '<strong>Right:</strong> Pairwise alignment of the consensus against the intended reference; highlighted regions mark positions where the two sequences do not match (mismatches or indels).'
        "</div>"
    )

    html_block = [
        '<div class="section-block">',
        "<h3>4. Consensus Sequence and Blast Result</h3>",
        '<div class="flex2">',
        left_box,
        right_box,
        "</div>",
        caption_html
    ]
    if warn_msgs:
        html_block.append(f'<div class="warning">{"；".join(warn_msgs)}</div>')
    html_block.append("</div>")
    return "\n".join(html_block)


def parse_contam_report(docx_path: Path):
    results = []
    if not HAS_DOCX or not docx_path.is_file():
        return results
    try:
        doc = Document(str(docx_path))
    except Exception:
        return results

    lines = [p.text.rstrip() for p in doc.paragraphs]
    n = len(lines)

    header_idx = [i for i, ln in enumerate(lines) if ln.strip().startswith(">umi")]
    if not header_idx:
        return results
    header_idx.append(n)

    contamination_id = 0
    for k in range(len(header_idx) - 1):
        start = header_idx[k]; end = header_idx[k + 1]
        block = lines[start:end]

        rate_val = None
        rate_rx = re.compile(r"The\s+contamination\s+rate\s+is\s+([0-9]*\.?[0-9]+)", re.IGNORECASE)
        for ln in block:
            m = rate_rx.search(ln)
            if m:
                try:
                    rate_val = float(m.group(1))
                except Exception:
                    rate_val = None
                break
        if rate_val is None or rate_val <= 0.005:
            continue

        key_rx = re.compile(r"consensus\s+sequence\s+of\s+contamination\s*[:：]\s*(.*)$", re.IGNORECASE)
        seq_collect = []
        for i, ln in enumerate(block):
            m = key_rx.search(ln)
            if m:
                tail = m.group(1).strip()
                if tail:
                    seq_collect.append(tail)
                for j in range(i + 1, len(block)):
                    ln2 = block[j].strip()
                    if not ln2 or ln2.startswith(">umi"):
                        break
                    seq_collect.append(ln2)
                break

        if seq_collect:
            raw = "".join(seq_collect)
            seq = "".join(re.findall(r"[ACGTUNacgtun]+", raw))
            if seq:
                contamination_id += 1
                results.append((f"contamination {contamination_id}", seq, rate_val))
    return results


def render_contam(sample):
    img_path = Path("contam_detect") / "result" / f"{sample}_contam_piechart.png"
    if img_path.is_file():
        pie_html = f"""
          <div class="figure">
            <img class="pie-small" src="{esc(str(img_path))}" alt="Contamination pie chart"/>
          </div>
        """
    else:
        pie_html = f'<div class="warning">找不到污染检测饼图：{esc(str(img_path))}</div>'

    report_path = Path("contam_detect") / "result" / f"{sample}common-polutionReport.docx"
    rows = parse_contam_report(report_path)

    if rows:
        table = [
            '<div class="tablewrap">',
            "<table>",
            "<thead><tr><th>Contamination</th><th>Sequence</th></tr></thead>",
            "<tbody>"
        ]
        for label, seq, rate in rows:
            table.append("<tr>")
            table.append(f"<td>{esc(label)}<br/>Contamination Rate: {rate}</td>")
            table.append(f'<td><div class="scrollbox">{esc(seq)}</div></td>')
            table.append("</tr>")
        table.append("</tbody></table></div>")
        table_html = "\n".join(table)
    else:
        table_html = ""  # 不显示 “No Contamination found”

    caption_html = (
        '<div class="caption tight">'
        "<strong>Contamination Detection</strong> — Pie chart summarizes contaminant types and their fractions among assigned reads (target vs. others). "
        "A positive contamination call is made only when a contaminant fraction exceeds 0.5%; remaining criteria follow our pipeline defaults: "
        "two-end signature concordance and alignment coverage ≤ 90% for the called contaminant. "
        "When contamination is detected, the report also displays per-contaminant consensus sequences (Racon×3 → Medaka) for verification."
        "</div>"
    )

    return f"""
    <div class="section-block">
      <h3>5. Contamination Detection</h3>
      {pie_html}
      {table_html}
      {caption_html}
    </div>
    """


def render_samples_header():
    return """
    <div class="section-block section-gap">
      <h2 id="samples" class="section-anchor">Ⅳ. Samples</h2>
    </div>
    """


def render_sample_section(sample, reads_map, stats_path):
    return f"""
    <div class="section-block">
      <h3 id="sample-{esc(sample)}" class="section-anchor">{esc(sample)}</h3>
      {render_reads_table(sample, reads_map, stats_path)}
      {render_igv_png_and_pdf(sample)}
      {render_blast(sample)}
      {render_contam(sample)}
    </div>
    """


# ------------------------------ 主流程 ------------------------------

def main():
    ap = argparse.ArgumentParser(description="Generate NPlasmid-seq HTML report (relative paths)")
    ap.add_argument("--samples", help="逗号分隔的样本名列表（与文件名一致）")
    ap.add_argument("--all", action="store_true", help="使用全体样本（自动识别）")
    ap.add_argument("--out", default="NPlasmid_seq_report.html", help="输出 HTML 文件名")
    args = ap.parse_args()

    stats_path = str(Path("Demultiplexed") / "Demultiplex_stats.txt")
    reads_map = read_demultiplex_stats(stats_path)

    if args.all:
        sample_list = sorted(set(list(reads_map.keys()) + [
            p.name.replace(".sorted_by_strand.png", "") for p in Path("IGV").glob("*.sorted_by_strand.png")
        ]))
    else:
        if not args.samples:
            print("请使用 --samples 或 --all 指定样本")
            return
        sample_list = [s.strip() for s in args.samples.split(",") if s.strip()]
    if not sample_list:
        print("未识别到样本名，请检查 --samples 参数或 Demultiplexed/IGV 目录")
        return

    parts = []
    parts.append(html_header())
    parts.append(html_nav(sample_list))
    parts.append(html_main_header())
    parts.append(section_intro())
    parts.append(section_workflow())
    parts.append(html_results_overview(sample_list))
    parts.append(render_samples_header())

    for s in sample_list:
        parts.append(render_sample_section(s, reads_map, stats_path))

    parts.append("""
  </div>
</main>
</div>
</body>
</html>
""")

    out = args.out
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(parts))
    print(f"[OK] Wrote: {out}")


if __name__ == "__main__":
    main()
