
# @Author: Mengting Hou, Yue Song, Yuyang Zhan
# @Time: 2023-07-11 17:10
# @File: CreateExcel.py



import sys
import os
import re
import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows

def complement(sequence):   # 保留
    # 辅助函数，返回序列的反向互补序列
    complement_dict = {'A': 'T', 'T': 'A', 'C': 'G', 'G': 'C'}
    return ''.join(complement_dict.get(base.upper(), base) for base in sequence[::-1])

def search_sgrna(file_path, sgrna):
    doc = Document(file_path)
    content = ''
    is_within_range = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.replace(' ', '')
        if text.lower() == "startofseq":
            is_within_range = True
        elif text.lower() == "endofseq":
            is_within_range = False
        elif is_within_range:
            content += text
    content = content.upper()
    if content.find(sgrna.upper().replace(' ', '')) != -1:
        return 'forward'
    elif content.find(complement(sgrna)) != -1:
        return 'reverse_complement'
    else:
        return None

def calculate_sequence_length(file_path):
    doc = Document(file_path)
    content = ''
    is_within_range = False
    found_start = False
    found_end = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.replace(' ', '')
        if text.lower() == "startofseq":
            is_within_range = True
            found_start = True
        elif text.lower() == "endofseq":
            is_within_range = False
            found_end = True
        elif is_within_range:
            content += text
    content = content.upper()
    sequence_length = len(content)

    if not found_start:
        print("Warning: 'startofseq' not found.")
    if not found_end:
        print("Warning: 'endofseq' not found.")
    if not found_start and not found_end:
        print("Warning: Neither 'startofseq' nor 'endofseq' found.")

    return sequence_length


# 获取文件夹路径
folder_path = os.path.dirname(os.path.abspath(__file__))   # 保留
# 获取目录下所有的docx文件
docx_files = [file for file in os.listdir(folder_path) if file.endswith('.docx')]

# 提取文件名中的信息并排序
data = []
for docx_file in docx_files:
    file_name = os.path.splitext(docx_file)[0]
    match = re.search(r'\s(.+)$', file_name)
    if match:
        if len(match.group(1)) > len(file_name.split()[0]):
            data.append((file_name.split()[0].strip(), match.group(1)))
        else:
            data.append((match.group(1), file_name.split()[0].strip()))


# 生成DataFrame
df = pd.DataFrame(data, columns=['plasmid-ID', 'plasmid-name'])

# 将第三到第五列内容置为空
df[['sgRNA-ID', 'sgRNA-seq', '备注']] = ''

# 拷贝"测序Cre文件.xlsx"的3-5列内容到"output.xlsx"
input_file = os.path.join(folder_path, '质粒-sgRNA-Pre.xlsx')   # 保留
output_file = os.path.join(folder_path, '质粒-sgRNA.xlsx')   # 保留
if os.path.exists(input_file):
    input_df = pd.read_excel(input_file)
    rows_to_copy = min(len(df), len(input_df))
    df.loc[:rows_to_copy, ['sgRNA-ID', 'sgRNA-seq', '备注']] = input_df.iloc[:rows_to_copy, 2:5].values
df.to_excel(output_file, index=False)

# 读取 Excel 文件
df = pd.read_excel(output_file)

df['Strand'] = ''
df['seq_length'] = ''
# 待匹配的 sgRNA 序列
sgrna_list = df['sgRNA-seq'].tolist()
sgrna_names = df['sgRNA-ID'].tolist()

for index, row in df.iterrows():
    plasmid_id = row["plasmid-ID"]
    found = False

    # 第一次尝试匹配 sgRNA 序列
    for sgrna, sgrna_name in zip(sgrna_list[:], sgrna_names[:]):
        for filename in os.listdir(folder_path): 
            if plasmid_id in filename:
                file_path = os.path.join(folder_path, filename)
                match_result = search_sgrna(file_path, sgrna)
                df.at[index, 'seq_length'] = calculate_sequence_length(file_path)
                if match_result == 'forward':
                    found = True
                    df.at[index, 'sgRNA-ID'] = sgrna_name
                    df.at[index, 'sgRNA-seq'] = sgrna
                    df.at[index, 'Strand'] = '+'
                    sgrna_list.remove(sgrna)
                    sgrna_names.remove(sgrna_name)
                    break
                elif match_result == 'reverse_complement':
                    found = True
                    df.at[index, 'sgRNA-ID'] = sgrna_name
                    df.at[index, 'sgRNA-seq'] = sgrna
                    df.at[index, 'Strand'] = '-'
                    sgrna_list.remove(sgrna)
                    sgrna_names.remove(sgrna_name)
                    break
        if found:
            break

    # 如果第一次匹配失败，回溯已匹配过的 sgRNA 进行再次尝试
    if not found:
        for sgrna, sgrna_name in zip(sgrna_list[:], sgrna_names[:]):
            reverse_complement = complement(sgrna)
            for filename in os.listdir(folder_path):  # 使用 folder_path 而不是 directory
                if plasmid_id in filename:
                    file_path = os.path.join(folder_path, filename)
                    match_result = search_sgrna(file_path, reverse_complement)
                    if match_result == 'reverse_complement':
                        found = True
                        df.at[index, 'sgRNA-ID'] = sgrna_name
                        df.at[index, 'sgRNA-seq'] = sgrna
                        df.at[index, 'Strand'] = '-'
                        sgrna_list.remove(sgrna)
                        sgrna_names.remove(sgrna_name)
                        break
        if not found:
            print(f"Warning: No match found for plasmid {plasmid_id}, please check sgRNA sequences!")

# 保存更新后的文件
df_sorted = df.sort_values(by="sgRNA-ID")
df_sorted.to_excel(output_file, index=False)