# -*- coding: utf-8 -*-
# @Time    : 2022/08/23 14:05
# @Author  : zzk  |  2025-09-01 全面修订：openpyxl/xls 支持、健壮性、报错提示
# @File    : blast_result1_5_only_blastresults.py
# Description: 从.fa读取共识序列，对每个参考建库+blast，比对结果写入docx

import os
import sys
import re
import glob
import time

# 文档/图片
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_COLOR_INDEX
# 可选：若后续启用 PDF->JPG，请先安装：pip install pdf2image
try:
    import pdf2image  # noqa: F401
except Exception:
    pdf2image = None

# Excel 读取：xlsx 用 openpyxl；xls 用 xlrd (2.0 起仅支持 .xls)
import openpyxl
import xlrd  # 仅用于 .xls

# -------------------------------
# 工具函数：读取 Excel（自动识别 .xlsx/.xls）
# -------------------------------
def load_excel_information():
    """
    搜索 ./ 目录下 Excel：优先包含 “质粒-sgRNA” 的文件；否则退化为任意 .xlsx/.xls。
    返回 information = [sample_time, Plasmid_id, ps, sgrna_id, sgrna_seq]
    """
    # cwd = os.getcwd()
    # # 先找带关键字
    # candidates = sorted(glob.glob("./*质粒-sgRNA*.xlsx")) + sorted(glob.glob("./*质粒-sgRNA*.xls"))
    # if not candidates:
    #     # 放宽：任意 Excel
    #     candidates = sorted(glob.glob("./*.xlsx")) + sorted(glob.glob("./*.xls"))

    # if not candidates:
    #     visible_xlsx = sorted(glob.glob("./*.xlsx"))
    #     visible_xls = sorted(glob.glob("./*.xls"))
    #     msg = [
    #         "未找到可用的 Excel 文件！请确认：",
    #         f"1) Excel 位于当前目录：{cwd}",
    #         "2) 建议文件名包含“质粒-sgRNA”（脚本优先匹配该关键词）",
    #         "3) 扩展名为 .xlsx 或 .xls",
    #         "",
    #         "当前目录下检测到的 Excel：",
    #         f"  .xlsx: {visible_xlsx if visible_xlsx else '无'}",
    #         f"  .xls : {visible_xls  if visible_xls  else '无'}",
    #     ]
    #     raise FileNotFoundError("\n".join(msg))
    cwd = os.getcwd()
    word_excel = os.path.join(os.path.dirname(cwd), "word", "质粒-sgRNA.xlsx")
    if os.path.exists(word_excel):
        candidates = [word_excel]
    else:
        candidates = sorted(glob.glob("./*质粒-sgRNA*.xlsx")) + sorted(glob.glob("./*质粒-sgRNA*.xls"))


    file_pre_name = candidates[0]
    ext = os.path.splitext(file_pre_name)[1].lower()

    Plasmid_id, ps, sgrna_id, sgrna_seq = [], [], [], []

    def _clean(x):
        if x is None:
            return ""
        return str(x).replace("\n", "").strip()

    if ext == ".xlsx":
        wb = openpyxl.load_workbook(file_pre_name, data_only=True)
        sh = wb.worksheets[0]
        rows = sh.max_row
        sample_time = _clean(sh.cell(row=1, column=1).value)
        for i in range(1, rows + 1):
            v1 = _clean(sh.cell(row=i, column=1).value)
            v2 = _clean(sh.cell(row=i, column=2).value)
            v3 = _clean(sh.cell(row=i, column=3).value)
            v4 = _clean(sh.cell(row=i, column=4).value)
            Plasmid_id.append(v1)
            ps.append(v2)
            sgrna_id.append(v3)
            sgrna_seq.append(v4)
    elif ext == ".xls":
        wb = xlrd.open_workbook(file_pre_name)
        sh = wb.sheets()[0]
        rows = sh.nrows
        sample_time = _clean(sh.row_values(0)[0] if rows > 0 else "")
        for i in range(rows):
            row = sh.row_values(i)
            c1 = _clean(row[0] if len(row) > 0 else "")
            c2 = _clean(row[1] if len(row) > 1 else "")
            c3 = _clean(row[2] if len(row) > 2 else "")
            c4 = _clean(row[3] if len(row) > 3 else "")
            Plasmid_id.append(c1)
            ps.append(c2)
            sgrna_id.append(c3)
            sgrna_seq.append(c4)
    else:
        raise ValueError(f"不支持的扩展名：{ext}（仅支持 .xlsx/.xls）")

    print("读取 Excel：", file_pre_name)
    print("样本时间(sample_time)：", sample_time)
    print("Plasmid_id 条数：", len(Plasmid_id))
    return [sample_time, Plasmid_id, ps, sgrna_id, sgrna_seq]


# -------------------------------
# 读取 .fa，拆分并保存到 ./consensus/，收集 name_number/consensus_num
# -------------------------------
def load_consensus_fa():
    # 优先使用 ../consensus/racon/consensus_racon.fa
    file_fa = os.path.join("..", "consensus", "racon", "consensus_racon.fa")
    if os.path.exists(file_fa):
        print(f"[INFO] 使用共识序列文件: {file_fa}")
    else:
        # fallback: 当前目录找 .fa
        candidates = [f for f in os.listdir(".") if f.endswith(".fa")]
        if not candidates:
            raise FileNotFoundError("未找到共识序列 .fa 文件！请确认 ../consensus/racon/consensus_racon.fa 或当前目录")
        file_fa = candidates[0]
        print(f"[INFO] 使用当前目录的共识序列文件: {file_fa}")

    if not os.path.exists("consensus"):
        os.makedirs("consensus")
    if not os.path.exists("doc_blast_result"):
        os.makedirs("doc_blast_result")

    pattern_2 = r"[P,S]{1}\d{3,4}[A,B]{0,1}"
    name_number, consensus_num = [], []

    with open(file_fa, "r") as f:
        content = f.readlines()
        for i in range(len(content) - 1):
            if content[i] and i % 2 == 0:
                # header 行：>xxx
                name_new = "_".join(content[i].replace("\n", "").split("_")[1:])
                m = re.search(pattern_2, name_new)
                if not m:
                    continue
                code = m.group(0)
                name_number.append(code)
                seq = content[i + 1].replace("\n", "")
                consensus_num.append(seq)
                with open(f"./consensus/{code}.fa", "w") as f1:
                    f1.write(content[i])
                    f1.write(content[i + 1])

    print("有共识序列的质粒有：", name_number)
    return name_number, consensus_num


# -------------------------------
# 生成 docx：把 BLAST 文本结果转成带高亮的 Word
# -------------------------------
def read_outfile(file, doc, information, name_number, consensus_num):
    """
    file: 比对结果文本文件名（与 makeblastdb/blastn 输出名一致）
    doc : docx.Document()
    information: [sample_time, Plasmid_id, ps, sgrna_id, sgrna_seq]
    """
    time_sample = information[0]
    if file not in information[1]:
        # 保护：Excel 不含此编号，跳过
        return "SKIP"

    index_in_p = information[1].index(file)
    ps_ = information[2][index_in_p]
    sgrna_id_ = information[3][index_in_p]
    sgrna_seq_ = information[4][index_in_p]

    doc.add_paragraph(time_sample)
    doc.add_paragraph(f"{file}     {ps_}")
    doc.add_paragraph(f"{sgrna_id_}     {sgrna_seq_}")
    doc.add_paragraph("  ")
    doc.add_paragraph("  ")

    # 写入 Consensus
    try:
        idx_cons = name_number.index(file)
        p_consensus0 = doc.add_paragraph()
        r0 = p_consensus0.add_run("Consensus:")
        r0.font.size = Pt(11); r0.bold = True

        p_consensus1 = doc.add_paragraph()
        r1 = p_consensus1.add_run(consensus_num[idx_cons])
        r1.font.name = "MS Mincho"
        r1.font.size = Pt(9)
    except ValueError:
        pass

    doc.add_paragraph("  ")
    doc.add_paragraph("BLASTN 2.9.0+")

    # 读取 BLAST 纯文本输出
    if not os.path.isfile(file):
        return "NO_FILE"

    with open(file, "r") as f:
        lines = f.readlines()
        # 若完全无命中
        for ln in lines:
            if "No hits found" in ln:
                return None

        i = 0
        while i < len(lines):
            if "Score" in lines[i][:11]:
                # 解析得分（原逻辑：>0 就输出）
                try:
                    score = float(lines[i].split(" ")[3])
                except Exception:
                    score = 0.0
                if score > 0:
                    doc.add_paragraph(f"{file}")
                    doc.add_paragraph("  ")
                    doc.add_paragraph("BLASTN 2.9.0+")
                    # 写该命中块
                    for ii in range(i, len(lines)):
                        if ii < len(lines):
                            if lines[ii] == "\n" and "Query" not in lines[ii + 1]:
                                end = ii
                                break
                        else:
                            end = ii
                            break

                        # 前4行加粗
                        if ii - i < 4:
                            p_score = doc.add_paragraph()
                            r = p_score.add_run(lines[ii].rstrip("\n"))
                            r.font.size = Pt(11); r.bold = True

                        elif (ii - i) % 4 == 0:
                            # 4行一个块：对齐高亮差异字符
                            # 这里参照原脚本，把第(ii+2)行与第(ii)行逐字符比对
                            count = (
                                lines[ii].upper().count("A")
                                + lines[ii].upper().count("T")
                                + lines[ii].upper().count("G")
                                + lines[ii].upper().count("C")
                                + lines[ii].upper().count("-")
                            )
                            # 行1（subject line）
                            p1 = doc.add_paragraph(); p1.paragraph_format.space_after = Pt(0)
                            r = p1.add_run(lines[ii + 2][:14]); r.font.name = "MS Mincho"; r.font.size = Pt(9)
                            for co in range(14, 14 + count):
                                ch_q = lines[ii + 2][co] if co < len(lines[ii + 2]) else " "
                                ch_s = lines[ii][co] if co < len(lines[ii]) else " "
                                if ch_q.upper() == ch_s.upper():
                                    r = p1.add_run(ch_q); r.font.name = "MS Mincho"; r.font.size = Pt(9)
                                else:
                                    r = p1.add_run(ch_q); r.font.name = "MS Mincho"; r.font.size = Pt(9)
                                    r.bold = True; r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            r = p1.add_run(lines[ii + 2][14 + count :].rstrip("\n")); r.font.name = "MS Mincho"; r.font.size = Pt(9)

                            # 行2（中间 ||| 行）
                            p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
                            r = p2.add_run(lines[ii + 1].rstrip("\n")); r.font.name = "MS Mincho"; r.font.size = Pt(9)

                            # 行3（query line）
                            p3 = doc.add_paragraph(); p3.paragraph_format.space_after = Pt(0)
                            r = p3.add_run(lines[ii][:14]); r.font.name = "MS Mincho"; r.font.size = Pt(9)
                            for co in range(14, 14 + count):
                                ch_q = lines[ii + 2][co] if co < len(lines[ii + 2]) else " "
                                ch_s = lines[ii][co] if co < len(lines[ii]) else " "
                                r = p3.add_run(ch_s); r.font.name = "MS Mincho"; r.font.size = Pt(9)
                                if ch_q.upper() != ch_s.upper():
                                    r.bold = True; r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            r = p3.add_run(lines[ii][14 + count :].rstrip("\n")); r.font.name = "MS Mincho"; r.font.size = Pt(9)

                            # 行4（坐标行）
                            p4 = doc.add_paragraph(); p4.paragraph_format.space_after = Pt(0)
                            r = p4.add_run(lines[ii + 3].rstrip("\n")); r.font.name = "MS Mincho"; r.font.size = Pt(9)

                        else:
                            continue

                    i = end + 1
                else:
                    i += 1
            else:
                i += 1

    # 保存 docx
    try:
        doc.save(f"./doc_blast_result/{file}.docx")
        return "Success"
    except Exception:
        print("保存失败，文件可能被占用，关闭重试！")
        sys.exit(1)


# -------------------------------
# 主流程
# -------------------------------
if __name__ == "__main__":
    # 新建 BLAST 工作目录
    base_dir = os.path.dirname(os.path.abspath(__file__))
    blast_dir = os.path.join(base_dir, "BLAST")
    os.makedirs(blast_dir, exist_ok=True)
    os.chdir(blast_dir)  # 切换到 BLAST 工作目录
    print(f"[INFO] BLAST 工作目录: {blast_dir}")

    # 1) 读取共识序列
    name_number, consensus_num = load_consensus_fa()

    # 2) 逐质粒建库 + BLAST
    pattern_2 = r"[P,S]{1}\d{3,4}[A,B]{0,1}"
    import subprocess

    for na in name_number:
        m = re.search(pattern_2, na)
        if not m:
            print(f"[WARN] 跳过无法解析的编号：{na}")
            continue

        # 建库
        cmd1 = f"makeblastdb -in ../Reference/{m.group(0)}.fasta -dbtype nucl -parse_seqids -out {na}"
        print(f"[CMD] {cmd1}")
        subprocess.run(cmd1, shell=True, check=True)

        # 比对
        cmd2 = f"blastn -query ./consensus/{na}.fa -db {na} -evalue 1e-6 -num_threads 6 -out {na}"
        print(f"[CMD] {cmd2}")
        subprocess.run(cmd2, shell=True, check=True)

    # 3) 读取 Excel 信息
    information = load_excel_information()

    # 4) 生成 docx（仅处理 Excel 中出现的编号，防止 index() 报错）
    for na in name_number:
        if na not in information[1]:
            print(f"[SKIP] {na} 不在 Excel 的 Plasmid_id 列中，已跳过。")
            continue
        print("file:", na)
        doc = docx.Document()
        status = read_outfile(na, doc, information, name_number, consensus_num)
        if status is None:
            print(f"[WARN] {na}: No hits found，未生成 docx。")
        elif status == "NO_FILE":
            print(f"[WARN] {na}: 未找到 BLAST 输出文件，跳过。")
        elif status == "Success":
            print(f"[OK] {na}: 已生成 ./doc_blast_result/{na}.docx")
        elif status == "SKIP":
            print(f"[SKIP] {na}: Excel 无此编号。")
